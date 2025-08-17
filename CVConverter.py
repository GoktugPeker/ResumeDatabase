import streamlit as st
from PyPDF2 import PdfReader
import sqlite3
from langchain_ollama import ChatOllama
from langchain.prompts import PromptTemplate,ChatPromptTemplate,MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import StreamlitChatMessageHistory
from docx import Document
import re
import io
import hashlib
from datetime import datetime, timezone, timedelta

CATEGORIES = [
    "Personal Information",
    "Professional Summary",
    "Work Experience",
    "Education",
    "Skills",
    "Certifications",
    "Projects",
    "References",
]

def split_into_sections(llm_text: str):
    """Split LLM output into your known CV sections (robust to **bold** or plain headings)."""
    text = (llm_text or "").replace("\r\n", "\n")
    pattern = re.compile(r"(?im)^\s*\**\s*(" + "|".join(re.escape(h) for h in CATEGORIES) + r")\s*\**\s*$")
    hits = list(pattern.finditer(text))
    sections = {c: "" for c in CATEGORIES}

    if not hits:
        sections["Professional Summary"] = text.strip()
        return sections

    for i, m in enumerate(hits):
        label = m.group(1)
        start = m.end()
        end = hits[i + 1].start() if i + 1 < len(hits) else len(text)
        body = text[start:end].strip()
        sections[label] = body

    return sections

st.title("CV Database with ChatBot using LLM")

if "saved_token" not in st.session_state:
    st.session_state["saved_token"] = None

# Storing Data in SQLite
conn = sqlite3.connect("cv_database.db")
cursor = conn.cursor()
cursor.execute("""
    CREATE TABLE IF NOT EXISTS cvs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT,
        content TEXT,
        uploaded_on DATETIME
    )
""")
conn.commit()

llm = ChatOllama(
    base_url="http://127.0.0.1:11434",
    model="qwen2.5:latest",
    temperature=0.1,
    num_predict=2000
)

topic = st.text_input("Ask something about CVs:")

text = ""

uploaded_file = st.file_uploader("Upload a CV  (pdf only)",
                                 accept_multiple_files=False,
                                 type=["pdf"])

if uploaded_file is not None:
    # Build a token for THIS selection only (does NOT block future re-uploads)
    file_bytes = uploaded_file.getvalue()
    selection_token = hashlib.md5(file_bytes).hexdigest() + f"|{uploaded_file.name}"

    # 1) Extract text once so BOTH branches can render the button
    text = ""
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            text += (page.extract_text() or "") + "\n"

    # 2) Insert only once per selection (debounce)
    if st.session_state["saved_token"] != selection_token:
        dubai = timezone(timedelta(hours=4))
        uploaded_on = datetime.now(dubai).strftime('%Y-%m-%d %H:%M:%S')

        cursor.execute(
            "INSERT INTO cvs (filename, content, uploaded_on) VALUES (?, ?, ?)",
            (uploaded_file.name, text, uploaded_on)
        )
        conn.commit()
        st.session_state["saved_token"] = selection_token
        st.success(f"Saved: {uploaded_file.name}")
    else:
        st.info("Already saved this selection. Skipping duplicate insert.")

else:
    # No file selected → next selection (even the same file) will insert again
    st.session_state["saved_token"] = None

# --- Build the template/chain ONCE (used for any selected CV) ---
### MOVED: template & chain outside upload block so they can be reused
cv_template = PromptTemplate(
    input_variables=["text"],
    template="""
You are an expert CV parser. Extract information ONLY from the provided CV text and output it in the EXACT schema and formatting below.

HARD BOUNDARY: Only use content strictly inside CV TEXT START ... CV TEXT END. If a fact is not present there, output "N/A". Do NOT use prior knowledge or assumptions.

CRITICAL RULES:
- Use EXACTLY these 8 headings (no others, no intro/outro text):
  Personal Information
  Professional Summary
  Work Experience
  Education
  Skills
  Certifications
  Projects
  References
- Under EACH heading, first output a single line starting with "Section Summary:" that gives a concise 1–2 sentence summary of the information in that section (or "N/A" if nothing is present).
- Then output the detailed items exactly as specified for that section.
- Never leave any section empty. If truly missing, write "N/A" for the section summary and any detailed fields.
- Do NOT fabricate or guess. Only use facts present in the CV text.
- Do NOT use code fences, tables, or extra markdown beyond the exact headings and hyphen bullets.
- Keep bullets as hyphens "- " (no numbering).
- Normalize dates as "Mon YYYY – Mon YYYY" or "Mon YYYY – Present"; if month missing, use "YYYY".
- Join multiple contact items with "; ".

MAPPING HINTS (non-exhaustive):
- Personal Information: name, email, phone, city/country, LinkedIn/GitHub/portfolio.
- Work Experience synonyms: Experience, Employment History, Professional Experience, Roles.
  • Project Name detection: look for explicit project/program/product/client/engagement names near each role. If multiple, pick the primary; if unclear, list up to 2 separated by "; ". If none, use "N/A".
  • Working Years: prefer date ranges attached to the role; if only years appear, use "YYYY – YYYY"; allow "Present".
  • Title: use the role title for that period; if multiple titles for overlapping dates, use the most specific/most recent per line.
  • Do NOT include employer, responsibilities, achievements, or locations in this section’s detailed lines.
- Education synonyms: Academics, Academic Background, Qualifications.
- Skills: technical and soft skills; deduplicate; keep concise.
- Certifications: certificates, licenses, badges.
- Projects: personal/work projects, case studies, publications framed as projects.

OUTPUT FORMAT (follow exactly; replace placeholders; keep headings identical; include N/A where unknown):

Personal Information
Section Summary: <1–2 sentences or N/A>
Full Name: <value or N/A>
Contact Information: <email; phone; location; other or N/A>
LinkedIn: <URL or N/A>

Professional Summary
Section Summary: <1–2 sentences or N/A>
<2–5 concise sentences expanding the summary, or N/A>

Work Experience
Section Summary: <1–2 sentences (e.g., total years, industries, key roles) or N/A>
<Title> – Project: <Project Name or N/A> | <Dates or N/A>
<Title> – Project: <Project Name or N/A> | <Dates or N/A>
(continue one line per role; no bullets; do NOT include employer/responsibilities/location; or write N/A)

Education
Section Summary: <1–2 sentences (e.g., highest degree, institutions, focus) or N/A>
<Degree> – <Institution>, <Location> | <Year or Dates or N/A>
<Another degree line if any>
(or N/A)

Skills
Section Summary: <1–2 sentences (e.g., core stack, domains) or N/A>
- <skill 1>
- <skill 2>
- <skill 3>
(up to ~15, or N/A)

Certifications
Section Summary: <1–2 sentences (e.g., notable certs) or N/A>
<Certification> – <Issuer> | <Year or N/A>
<Another certification if any>
(or N/A)

Projects
Section Summary: <1–2 sentences (e.g., notable projects, impact) or N/A>
<Project Title> – <brief description 1 line>
<Another project line if any>
(or N/A)

References
Section Summary: <1–2 sentences (e.g., provided vs upon request) or N/A>
<If explicitly provided, list; otherwise "References available upon request">

CV TEXT START
{text}
CV TEXT END
"""
)
chain_cv = cv_template | llm

# --- NEW: Dropdown to pick any saved CV and generate DOCX on demand ---
cursor.execute("SELECT id, filename, content, uploaded_on FROM cvs")
all_rows = cursor.fetchall()
options = [
    {"id": r[0], "filename": r[1], "content": r[2], "uploaded_on": r[3]}
    for r in all_rows
]

selected = st.selectbox(
    "Choose a saved CV to format (type to search):",
    options,
    index=None,
    format_func=lambda o: f"{o['filename']} — {o['uploaded_on']}" if o else "—",
    placeholder="Start typing a filename…"
)

if selected:
    if st.button("Generate DOCX for selected CV"):
        selected_text = selected["content"] or ""
        if selected_text.strip():
            formatted_output = chain_cv.invoke({"text": selected_text})
            result_text = getattr(formatted_output, "content", str(formatted_output))

            sections = split_into_sections(result_text)

            doc = Document()
            doc.add_heading('Curriculum Vitae', level=0)
            for cat in CATEGORIES:
                doc.add_heading(cat, level=1)
                body = sections.get(cat, "").strip()
                if body:
                    for line in body.splitlines():
                        if line.strip():
                            doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")

            docx_buf = io.BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)

            base_name = (selected["filename"] or "cv").rsplit(".", 1)[0]
            st.download_button(
                label="⬇️ Download DOCX (templated)",
                data=docx_buf,
                file_name=f"{base_name}_formatted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.info("Selected CV has no extracted text.")

# Build a combined corpus for QA if no fresh `text` from current upload
if not text:
    cursor.execute("SELECT id, filename, content FROM cvs")
    rows = cursor.fetchall()
    text = "\n\n".join([f"{r[1] or 'unknown'}:\n{r[2] or ''}" for r in rows])

# --- QA prompt/chain available regardless of upload ---
qa_prompt = ChatPromptTemplate.from_messages([
    ("system",
        "Answer the user's question by checking all uploaded documents."
        "Use any data uploaded"
        "Fetch all the CVs and data in CV text give best possible answer "
        "If the answer is not contained in the CV, reply exactly: I don't know.\n\n"
        "CV Text:\n{text}\n\n"
        "Question: {topic}"
    ),
    MessagesPlaceholder(variable_name="chat_history"),
    ("human","{topic}")
])
chain_qa = qa_prompt | llm 

history_for_chain = StreamlitChatMessageHistory(key="chat_history")
chain_with_history = RunnableWithMessageHistory(
    chain_qa,
    lambda session_id: history_for_chain,
    input_messages_key="topic",
    history_messages_key="chat_history"
)

# --- Run chatbot answer if a question is provided ---
if topic:
    if not text.strip():
        st.info("No CVs available in the database to answer from.")
    else:
        st.subheader("💬 Chatbot answer")

        # Option A: simple write_stream
        def to_text(gen):
            for chunk in gen:
                # chunk is an AIMessageChunk (preferred) or occasionally a str
                yield (getattr(chunk, "content", None) 
                       or (chunk if isinstance(chunk, str) else ""))

        st.write_stream(to_text(
            chain_with_history.stream(
                {'text': text, 'topic': topic},
                {"configurable": {'session_id': 'abc0002'}}
            )
        ))

with st.expander("Conversation History"):
    msgs = history_for_chain.messages  # list[BaseMessage]
    if not msgs:
        st.caption("No messages yet.")
    else:
        for m in msgs:
            role = m.type.title()  # "Human" / "AI" / "System"
            st.markdown(f"**{role}:** {m.content}")

# optional: a clear button
if st.button("Clear chat history"):
    history_for_chain.clear()
    st.success("Chat history cleared.")

conn.close()
